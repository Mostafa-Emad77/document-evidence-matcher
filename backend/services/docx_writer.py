"""
DOCX writer — generates a source-like review document.

Output behavior:
  - Preserve extracted paragraph order from the source content
  - Highlight quote paragraphs for fast visual scanning
  - Add internal links from each quote back to its original description
  - Embed matched screenshot images below each quote block
"""
from __future__ import annotations

import base64
import io
from collections import defaultdict
from typing import Any

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from backend.schemas import ElementType, Milestone1Output


def _norm(text: str) -> str:
    return " ".join(text.split()).strip().lower()


def _add_bookmark(paragraph, bookmark_name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))

    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_internal_hyperlink(paragraph, anchor: str, text: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")

    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)

    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)

    paragraph._p.append(hyperlink)


def _split_quote_parts_with_offsets(quote_text: str) -> list[tuple[str, int, int]]:
    """
    Split joined quote text back into per-paragraph parts and preserve offsets.
    """
    parts: list[tuple[str, int, int]] = []
    cursor = 0
    for raw in quote_text.split("\n\n"):
        start = cursor
        end = start + len(raw)
        cursor = end + 2
        stripped = raw.strip()
        if not stripped:
            continue
        lead_trim = len(raw) - len(raw.lstrip())
        trail_trim = len(raw) - len(raw.rstrip())
        part_start = start + lead_trim
        part_end = end - trail_trim
        if part_start < part_end:
            parts.append((stripped, part_start, part_end))
    return parts


def _enum_color(color_name: str, default: WD_COLOR_INDEX | None) -> WD_COLOR_INDEX | None:
    return getattr(WD_COLOR_INDEX, color_name, default)


def _normalize_with_map(text: str) -> tuple[str, list[int]]:
    out_chars: list[str] = []
    out_map: list[int] = []
    prev_space = False

    for idx, ch in enumerate(text):
        if ch.isspace():
            if not out_chars or prev_space:
                continue
            out_chars.append(" ")
            out_map.append(idx)
            prev_space = True
            continue
        out_chars.append(ch)
        out_map.append(idx)
        prev_space = False

    if out_chars and out_chars[-1] == " ":
        out_chars.pop()
        out_map.pop()
    return "".join(out_chars), out_map


def _project_spans_to_text(
    source_text: str,
    target_text: str,
    spans: list[tuple[int, int, str]],
) -> list[tuple[int, int, str]]:
    """
    Re-map source-text offsets into target-text offsets after whitespace drift.
    """
    if not spans:
        return []
    if source_text == target_text:
        return spans

    target_norm, target_map = _normalize_with_map(target_text)
    if not target_norm or not target_map:
        return []

    used_target: list[tuple[int, int]] = []
    projected: list[tuple[int, int, str]] = []
    target_norm_cursor = 0

    for start, end, color in sorted(spans, key=lambda item: item[0]):
        if end <= start:
            continue
        if start < 0 or end > len(source_text):
            continue

        src_slice = source_text[start:end]
        normalized_slice = " ".join(src_slice.split()).strip()
        if not normalized_slice:
            continue

        norm_start = target_norm.find(normalized_slice, target_norm_cursor)
        if norm_start < 0:
            norm_start = target_norm.find(normalized_slice)
        if norm_start < 0:
            continue
        norm_end = norm_start + len(normalized_slice)

        tgt_start = target_map[norm_start]
        tgt_end = target_map[norm_end - 1] + 1
        overlap = any(not (tgt_end <= a or tgt_start >= b) for a, b in used_target)
        if overlap:
            continue

        projected.append((tgt_start, tgt_end, color))
        used_target.append((tgt_start, tgt_end))
        target_norm_cursor = norm_end

    return projected


def _build_local_spans(
    spans: list[Any],
    side: str,
    part_start: int | None = None,
    part_end: int | None = None,
) -> list[tuple[int, int, str]]:
    out: list[tuple[int, int, str]] = []
    for span in spans:
        if side == "description":
            start = int(span.description_start)
            end = int(span.description_end)
        else:
            start = int(span.quote_start)
            end = int(span.quote_end)
        if end <= start:
            continue
        if part_start is not None and part_end is not None:
            # Keep the in-part intersection so cross-paragraph matches still
            # highlight the visible overlap in this part.
            start = max(start, part_start)
            end = min(end, part_end)
            if end <= start:
                continue
            start -= part_start
            end -= part_start
        out.append((start, end, str(span.color)))
    unique: list[tuple[int, int, str]] = []
    seen: set[tuple[int, int]] = set()
    for start, end, color in sorted(out, key=lambda item: item[0]):
        key = (start, end)
        if key in seen:
            continue
        seen.add(key)
        unique.append((start, end, color))
    return unique


def _add_colored_runs(
    paragraph,
    text: str,
    spans: list[tuple[int, int, str]],
    default_highlight: WD_COLOR_INDEX | None,
) -> None:
    if not spans:
        run = paragraph.add_run(text)
        run.font.highlight_color = default_highlight
        return

    cursor = 0
    text_len = len(text)
    # Sort by start so we emit text left-to-right; spans may come from multiple
    # sources (e.g. a paragraph's own link plus a borrowed-highlight anchor)
    # and could overlap, so we clip each span against the cursor and drop any
    # remainder that has already been emitted.
    ordered = sorted(spans, key=lambda item: (item[0], item[1]))
    for start, end, color_name in ordered:
        start = max(0, min(start, text_len))
        end = max(0, min(end, text_len))
        if end <= start:
            continue
        if start < cursor:
            start = cursor
        if end <= start:
            continue

        if start > cursor:
            gap_run = paragraph.add_run(text[cursor:start])
            gap_run.font.highlight_color = default_highlight

        color_run = paragraph.add_run(text[start:end])
        color_run.font.highlight_color = _enum_color(color_name, default_highlight)
        cursor = end

    if cursor < text_len:
        trail_run = paragraph.add_run(text[cursor:])
        trail_run.font.highlight_color = default_highlight


def _find_best_quote_plan_entry(
    quote_text_norm: str,
    quote_plan: list[dict[str, Any]],
    start_index: int,
) -> tuple[dict[str, Any] | None, int]:
    """
    Find the best quote plan entry for a quote segment using ordered matching.
    Returns (entry, next_plan_index).
    """
    if not quote_text_norm:
        return None, start_index

    for idx in range(start_index, len(quote_plan)):
        part_norm = quote_plan[idx]["part_norm"]
        if quote_text_norm == part_norm or quote_text_norm in part_norm or part_norm in quote_text_norm:
            return quote_plan[idx], idx + 1

    # Fallback search in case the sequence drifted for a specific document.
    for idx in range(0, start_index):
        part_norm = quote_plan[idx]["part_norm"]
        if quote_text_norm == part_norm or quote_text_norm in part_norm or part_norm in quote_text_norm:
            return quote_plan[idx], start_index

    return None, start_index


def _segment_position(match: Any) -> int | None:
    seg_id = str(getattr(match, "segment_id", "") or "")
    if not seg_id.startswith("seg_"):
        return None
    try:
        return int(seg_id.split("_", 1)[1])
    except (TypeError, ValueError):
        return None


def _is_last_quote_segment(segments: list[Any], current_index: int) -> bool:
    for next_match in segments[current_index + 1:]:
        next_text = (getattr(next_match, "segment_text", "") or "").strip()
        if not next_text:
            continue
        return getattr(next_match, "segment_type", None) != ElementType.quote
    return True


def _maybe_embed_matched_screenshots(doc: Document, match: Any) -> None:
    images = list(getattr(match, "matched_images", []) or [])
    if not images:
        return

    unique_images: list[Any] = []
    seen_keys: set[str] = set()
    for image in images:
        img_key = str(getattr(image, "image_id", "") or getattr(image, "thumbnail_base64", "") or "")
        if not img_key or img_key in seen_keys:
            continue
        seen_keys.add(img_key)
        unique_images.append(image)

    if not unique_images:
        return

    total = len(unique_images)
    for idx, image in enumerate(unique_images, start=1):
        img_b64 = getattr(image, "thumbnail_base64", None)
        if not img_b64:
            continue
        try:
            img_bytes = base64.b64decode(img_b64)
            img_stream = io.BytesIO(img_bytes)
            img_para = doc.add_paragraph()
            img_run = img_para.add_run()
            img_run.add_picture(img_stream, width=Inches(4.5))
            if total > 1:
                caption_text = f"↑ Matched screenshot evidence ({idx}/{total})"
            else:
                caption_text = "↑ Matched screenshot evidence"
            caption_para = doc.add_paragraph(caption_text)
            caption_run = caption_para.runs[0]
            caption_run.font.size = Pt(8)
            caption_run.font.italic = True
            caption_run.font.color.rgb = None  # inherit theme colour
        except Exception:
            # Skip malformed images but keep rendering remaining evidence.
            continue


def generate_docx(output: Milestone1Output) -> bytes:
    """
    Build a source-like DOCX with highlighted quotes and quote->description links.
    """
    doc = Document()

    # Styles
    styles = doc.styles
    if "Quote Block" not in [s.name for s in styles]:
        qs = styles.add_style("Quote Block", 1)  # 1 = paragraph style
        qs.base_style = styles["Normal"]
        qs.font.italic = True
        pfmt = qs.paragraph_format
        pfmt.left_indent = 0
        pfmt.right_indent = 0

    quote_plan: list[dict[str, Any]] = []
    quote_plan_by_position: dict[int, dict[str, Any]] = {}
    desc_to_anchors_by_norm: dict[str, list[str]] = defaultdict(list)
    desc_to_anchors_by_position: dict[int, list[str]] = defaultdict(list)
    anchor_to_desc_data: dict[str, dict[str, Any]] = {}
    for index, ql in enumerate(output.quote_table, start=1):
        anchor = f"desc_{index}"
        desc_norm = _norm(ql.description_text)
        desc_to_anchors_by_norm[desc_norm].append(anchor)
        if ql.description_position is not None:
            desc_to_anchors_by_position[int(ql.description_position)].append(anchor)

        description_side_spans = _build_local_spans(ql.claim_spans, "description")
        borrowed = bool(ql.colorization_description_text)
        # When highlights are borrowed from a previous description, the span
        # indices reference that borrowed text. Keep the link's own anchor
        # bookmark on the short intro paragraph (so "Go to description" still
        # works), but do NOT attach the description-side spans to it — they
        # belong on the borrowed paragraph instead.
        anchor_to_desc_data[anchor] = {
            "source_text": ql.description_text,
            "spans": [] if borrowed else description_side_spans,
        }

        if borrowed:
            borrow_anchor = f"desc_{index}_borrow"
            borrow_norm = _norm(ql.colorization_description_text or "")
            if borrow_norm:
                desc_to_anchors_by_norm[borrow_norm].append(borrow_anchor)
            if ql.colorization_description_position is not None:
                desc_to_anchors_by_position[
                    int(ql.colorization_description_position)
                ].append(borrow_anchor)
            anchor_to_desc_data[borrow_anchor] = {
                "source_text": ql.colorization_description_text or "",
                "spans": description_side_spans,
                "is_borrow": True,
            }

        quote_start_position = int(ql.quote_start_position) if ql.quote_start_position is not None else None
        quote_end_position = int(ql.quote_end_position) if ql.quote_end_position is not None else None
        quote_parts = _split_quote_parts_with_offsets(ql.quote_text)
        for part_idx, (part_text, part_start, part_end) in enumerate(quote_parts):
            part_entry = {
                "part_norm": _norm(part_text),
                "anchor": anchor,
                "source_text": part_text,
                "spans": _build_local_spans(ql.claim_spans, "quote", part_start, part_end),
                "citation_text": ql.citation_text,
                "is_last_part": part_idx == (len(quote_parts) - 1),
            }
            quote_plan.append(part_entry)
            if quote_start_position is None or quote_end_position is None:
                continue
            quote_position = quote_start_position + part_idx
            if quote_position > quote_end_position:
                continue
            quote_plan_by_position[quote_position] = part_entry

    remaining_desc_anchors: dict[str, list[str]] = {
        desc: list(anchors)
        for desc, anchors in desc_to_anchors_by_norm.items()
    }
    remaining_desc_anchors_by_position: dict[int, list[str]] = {
        pos: list(anchors)
        for pos, anchors in desc_to_anchors_by_position.items()
    }
    created_bookmarks: set[str] = set()
    quote_plan_index = 0
    next_bookmark_id = 1

    for seg_idx, match in enumerate(output.segments):
        text = (match.segment_text or "").strip()
        if not text:
            continue

        if match.segment_type == ElementType.heading:
            doc.add_paragraph(text, style="Heading 2")
            continue

        if match.segment_type == ElementType.narration:
            paragraph = doc.add_paragraph(style="Normal")
            seg_position = _segment_position(match)
            desc_norm = _norm(text)
            anchors = (
                remaining_desc_anchors_by_position.get(seg_position, [])
                if seg_position is not None
                else []
            )
            if not anchors:
                anchors = remaining_desc_anchors.get(desc_norm, [])
            while anchors and anchors[0] in created_bookmarks:
                anchors.pop(0)

            # Collect projected spans from EVERY pending anchor at this paragraph
            # so that borrowed highlights (from a later quote whose own
            # description was too short) render alongside the paragraph's own
            # highlights.
            merged_projected_spans: list[tuple[int, int, str]] = []
            for pending_anchor in anchors:
                desc_data = anchor_to_desc_data.get(pending_anchor, {})
                projected = _project_spans_to_text(
                    str(desc_data.get("source_text", "")),
                    text,
                    list(desc_data.get("spans", [])),
                )
                if projected:
                    merged_projected_spans.extend(projected)

            if merged_projected_spans:
                _add_colored_runs(
                    paragraph,
                    text,
                    merged_projected_spans,
                    default_highlight=None,
                )
            else:
                paragraph.add_run(text)

            # Add bookmarks for every pending anchor mapped to this description.
            while anchors:
                anchor = anchors.pop(0)
                _add_bookmark(paragraph, anchor, next_bookmark_id)
                created_bookmarks.add(anchor)
                next_bookmark_id += 1

        elif match.segment_type == ElementType.quote:
            paragraph = doc.add_paragraph(style="Quote Block")
            seg_position = _segment_position(match)
            plan_entry = quote_plan_by_position.get(seg_position) if seg_position is not None else None
            if plan_entry is None:
                plan_entry, quote_plan_index = _find_best_quote_plan_entry(_norm(text), quote_plan, quote_plan_index)
            if plan_entry:
                projected_quote_spans = _project_spans_to_text(
                    str(plan_entry.get("source_text", "")),
                    text,
                    list(plan_entry.get("spans", [])),
                )
                if projected_quote_spans:
                    _add_colored_runs(
                        paragraph,
                        text,
                        projected_quote_spans,
                        default_highlight=None,
                    )
                else:
                    paragraph.add_run(text)
            else:
                paragraph.add_run(text)

            anchor = plan_entry["anchor"] if plan_entry else None
            if anchor and anchor in created_bookmarks:
                paragraph.add_run(" ")
                _add_internal_hyperlink(paragraph, anchor, "[Go to description]")

            should_emit_citation = bool(
                plan_entry.get("is_last_part", True)
                if plan_entry is not None
                else _is_last_quote_segment(output.segments, seg_idx)
            )
            citation_text = str(plan_entry.get("citation_text", "")).strip() if plan_entry else ""
            if not citation_text:
                citation_text = (match.matched_citation_text or "").strip()
            if should_emit_citation and citation_text:
                citation_para = doc.add_paragraph(citation_text)
                if citation_para.runs:
                    citation_run = citation_para.runs[0]
                    citation_run.font.size = Pt(8)
                    citation_run.font.italic = True
                    citation_run.font.color.rgb = None  # inherit theme colour

            if should_emit_citation:
                _maybe_embed_matched_screenshots(doc, match)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
